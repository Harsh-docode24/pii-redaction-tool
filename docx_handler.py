"""DOCX Handler - Read and write .docx files preserving formatting."""
import re
from typing import List, Tuple
from docx import Document


class DocxHandler:
    """Handles reading and writing DOCX files with formatting preservation."""

    def __init__(self, input_path: str):
        """Load a DOCX file."""
        self.input_path = input_path
        self.doc = Document(input_path)

    def get_full_text(self) -> str:
        """Extract all text from the document (paragraphs + tables)."""
        full_text = []
        for para in self.doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return '\n'.join(full_text)

    def redact_and_save(self, replacements: List[Tuple[str, str]], output_path: str):
        """Apply replacements to the document and save.
        
        Args:
            replacements: List of (original, replacement) tuples
            output_path: Path to save the redacted document
        """
        # Sort replacements by length (longest first) to avoid partial matches
        replacements_sorted = sorted(replacements, key=lambda x: len(x[0]), reverse=True)
        
        # Apply to paragraphs
        for para in self.doc.paragraphs:
            self._replace_in_paragraph(para, replacements_sorted)

        # Apply to tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements_sorted)

        # Apply to headers/footers
        for section in self.doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header and header.is_linked_to_previous is False:
                    for para in header.paragraphs:
                        self._replace_in_paragraph(para, replacements_sorted)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and footer.is_linked_to_previous is False:
                    for para in footer.paragraphs:
                        self._replace_in_paragraph(para, replacements_sorted)

        self.doc.save(output_path)

        # Post-save: scrub emails from hidden XML (instrText hyperlinks, rels)
        self._scrub_xml_hyperlinks(output_path, replacements_sorted)

    def _scrub_xml_hyperlinks(self, docx_path: str, replacements: List[Tuple[str, str]]):
        """Replace PII inside DOCX XML that python-docx doesn't expose.
        
        This handles:
        - w:instrText containing HYPERLINK "mailto:original@email.com"
        - Relationship targets in .rels files
        - Any remaining mailto: links from known PII domains
        """
        import zipfile
        import os
        from faker import Faker
        fake = Faker()
        Faker.seed(42)

        # Build email-only replacements for XML scrubbing
        email_replacements = [(orig, repl) for orig, repl in replacements
                              if '@' in orig]

        # Known PII domains - any mailto: to these domains must be scrubbed
        pii_domains = [
            'kshinternational.com', 'kshinterantional.com', 'nuvama.com',
            'icicisecurities.com', 'hdfcbank.com', 'icicibank.com',
            'trilegal.com', 'in.mpms.mufg.com', 'kirtanepandit.com',
            'citi.com', 'eximbankindia.in', 'indusind.com',
            'sbi.co.in', 'federalbank.co.in', 'bajajfinserv.in',
        ]

        # Read the DOCX (which is a ZIP), modify XML in memory, rewrite
        temp_path = docx_path + '.tmp'
        with zipfile.ZipFile(docx_path, 'r') as zin:
            with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    data = zin.read(item)

                    # Only process XML and rels files
                    if item.endswith('.xml') or item.endswith('.rels'):
                        text = data.decode('utf-8')
                        
                        # Pass 1: Replace known detected emails
                        for orig_email, fake_email in email_replacements:
                            text = text.replace(f'mailto:{orig_email}', f'mailto:{fake_email}')
                            text = text.replace(orig_email, fake_email)
                            # Case-insensitive variant
                            if orig_email.lower() != orig_email:
                                text = re.compile(re.escape(orig_email), re.IGNORECASE).sub(
                                    fake_email, text)

                        # Pass 2: Sweep for ANY remaining emails from PII domains
                        remaining = re.findall(
                            r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}', text)
                        for email_addr in remaining:
                            domain = email_addr.split('@')[1].lower()
                            if any(d == domain or domain.endswith('.' + d) for d in pii_domains):
                                fake_local = fake.user_name()
                                fake_replacement = f"{fake_local}@redacted.example.com"
                                text = text.replace(f'mailto:{email_addr}', f'mailto:{fake_replacement}')
                                text = text.replace(email_addr, fake_replacement)

                        data = text.encode('utf-8')

                    zout.writestr(item, data)

        # Replace original with scrubbed version
        os.replace(temp_path, docx_path)

    def _replace_in_paragraph(self, paragraph, replacements: List[Tuple[str, str]]):
        """Replace PII in a paragraph while preserving formatting."""
        full_text = paragraph.text
        if not full_text.strip():
            return

        for original, replacement in replacements:
            if original.lower() in full_text.lower():
                self._replace_in_runs(paragraph, original, replacement)
                full_text = paragraph.text

    def _replace_in_runs(self, paragraph, search_text: str, replace_text: str):
        """Replace text across runs, handling cases where text spans multiple runs."""
        # First try simple run-by-run replacement
        for run in paragraph.runs:
            if search_text.lower() in run.text.lower():
                run.text = self._case_insensitive_replace(run.text, search_text, replace_text)
                return

        # If the text spans multiple runs, we need a cross-run approach
        runs = paragraph.runs
        if not runs:
            return

        combined_text = ''.join(r.text for r in runs)
        lower_combined = combined_text.lower()
        lower_search = search_text.lower()
        
        start_idx = lower_combined.find(lower_search)
        if start_idx == -1:
            return

        end_idx = start_idx + len(search_text)
        
        # Find which runs the match spans
        char_count = 0
        start_run = None
        end_run = None
        start_offset = 0
        end_offset = 0
        
        for i, run in enumerate(runs):
            run_start = char_count
            run_end = char_count + len(run.text)
            
            if start_run is None and run_end > start_idx:
                start_run = i
                start_offset = start_idx - run_start
            
            if run_end >= end_idx:
                end_run = i
                end_offset = end_idx - run_start
                break
            
            char_count = run_end

        if start_run is None or end_run is None:
            return

        if start_run == end_run:
            run = runs[start_run]
            run.text = run.text[:start_offset] + replace_text + run.text[end_offset:]
        else:
            runs[start_run].text = runs[start_run].text[:start_offset] + replace_text
            for i in range(start_run + 1, end_run):
                runs[i].text = ''
            runs[end_run].text = runs[end_run].text[end_offset:]

    @staticmethod
    def _case_insensitive_replace(text: str, search: str, replacement: str) -> str:
        """Replace text case-insensitively."""
        pattern = re.compile(re.escape(search), re.IGNORECASE)
        return pattern.sub(replacement, text)
